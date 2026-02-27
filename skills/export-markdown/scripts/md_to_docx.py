#!/usr/bin/env python3
"""Convert markdown to Word (.docx). Generic converter for standard markdown.

Usage:
  python md_to_docx.py input.md [output.docx]
  python md_to_docx.py input.md              # writes input.docx

Requires: pip install python-docx
"""
from docx import Document
import re
import sys
from pathlib import Path


def add_bold_para(doc, label, rest=''):
    """Add paragraph with bold label and optional rest."""
    p = doc.add_paragraph()
    r1 = p.add_run(label + (': ' if rest else ''))
    r1.bold = True
    if rest:
        p.add_run(rest)


def add_inline_bold_para(doc, text):
    """Add paragraph, rendering **bold** as bold."""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)


def is_table_separator(line):
    """Check if line is |---|---| style."""
    stripped = line.strip()
    if not stripped.startswith('|'):
        return False
    return bool(re.match(r'^\|[\s\-:]+\|$', stripped))


def md_to_docx(md_path, docx_path):
    doc = Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    table_rows = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)

        # Horizontal rule
        elif stripped == '---':
            doc.add_paragraph()

        # Table row
        elif stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            if not is_table_separator(line):
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                if any(c for c in cells):
                    table_rows.append(cells)
        else:
            if in_table and table_rows:
                num_cols = max(len(r) for r in table_rows)
                t = doc.add_table(rows=len(table_rows), cols=num_cols)
                t.style = 'Table Grid'
                for ri, row in enumerate(table_rows):
                    for ci in range(num_cols):
                        cell_text = row[ci] if ci < len(row) else ''
                        t.rows[ri].cells[ci].text = cell_text
                table_rows = []
                in_table = False

            if not stripped:
                doc.add_paragraph()
            elif re.match(r'^\*\*[^*]+\*\*:\s*.+', stripped):
                m = re.match(r'^\*\*([^*]+)\*\*:\s*(.+)', stripped)
                if m:
                    add_bold_para(doc, m.group(1), m.group(2))
            elif re.match(r'^\*\*[^*]+\*\*\s*$', stripped):
                add_bold_para(doc, stripped[2:-2].strip())
            elif stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped and not stripped.startswith('#'):
                add_inline_bold_para(doc, stripped)

        i += 1

    if table_rows:
        num_cols = max(len(r) for r in table_rows)
        t = doc.add_table(rows=len(table_rows), cols=num_cols)
        t.style = 'Table Grid'
        for ri, row in enumerate(table_rows):
            for ci in range(num_cols):
                cell_text = row[ci] if ci < len(row) else ''
                t.rows[ri].cells[ci].text = cell_text

    doc.save(docx_path)
    return docx_path


def main():
    if len(sys.argv) < 2:
        print(__doc__.strip())
        print("\nUsage: python md_to_docx.py input.md [output.docx]")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: {md_path} not found")
        sys.exit(1)

    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2])
    else:
        docx_path = md_path.with_suffix('.docx')

    out = md_to_docx(str(md_path), str(docx_path))
    print(f"Saved: {out}")


if __name__ == '__main__':
    main()
